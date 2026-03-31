import os
import json
import whisper
import sounddevice as sd
import numpy as np
from scipy.io.wavfile import write
from fpdf import FPDF
from docx import Document
import pdfplumber
import PyPDF2
import re
import random
import shutil
import glob
from collections import Counter
from datetime import datetime

try:
    import requests
    HAS_OLLAMA = True
except ImportError:
    HAS_OLLAMA = False

try:
    import pytesseract
    from PIL import Image
    import io
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


samplerate = 44100
_WHISPER_MODEL = None
OLLAMA_BASE_URL = "http://localhost:11434"
OLLAMA_API_URL = f"{OLLAMA_BASE_URL}/api/generate"
OLLAMA_MODEL = "mistral"  # Fast & capable model


# ============================================================
# OLLAMA SETUP (Local AI - Free & Unlimited)
# ============================================================
def get_ollama_client():
    """Check if Ollama is running."""
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=2)
        response.raise_for_status()
        print("🧠 Ollama connected (using local models - free & unlimited)")
        return True
    except Exception as e:
        print(f"⚠ Ollama not running. Please install Ollama from https://ollama.ai and run: ollama serve")
        print(f"   Then in another terminal: ollama pull {OLLAMA_MODEL}")
        print(f"   Error: {e}")
        return False


def ollama_generate(prompt, temperature=0.3, max_tokens=4096):
    """Call Ollama locally and return text response."""
    try:
        payload = {
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "options": {
                "temperature": temperature,
                "top_p": 0.95,
                "num_predict": max(128, min(int(max_tokens), 1200)),
            },
            "stream": False,
        }
        response = requests.post(OLLAMA_API_URL, json=payload, timeout=60)
        if response.status_code == 200:
            result = response.json()
            return result.get("response", "").strip()
        else:
            print(f"⚠ Ollama API error: {response.status_code} - {response.text}")
            return ""
    except Exception as e:
        print(f"⚠ Ollama connection error: {e}")
        return ""


def ollama_generate_json(prompt, temperature=0.2, max_tokens=4096):
    """Call Ollama and parse JSON response."""
    full_prompt = prompt + "\n\nIMPORTANT: Respond ONLY with valid JSON. No markdown fences, no preamble, no explanation."
    raw = ollama_generate(full_prompt, temperature, max_tokens)
    # Strip markdown fences if present
    raw = re.sub(r'^```(?:json)?\s*', '', raw)
    raw = re.sub(r'\s*```$', '', raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"⚠ JSON parse error: {e}\nRaw response (first 500 chars): {raw[:500]}")
        return None


# ============================================================
# WHISPER SETUP
# ============================================================
def get_whisper_model():
    global _WHISPER_MODEL
    if _WHISPER_MODEL is None:
        model_name = os.getenv("WHISPER_MODEL", "base")
        print(f"🧠 Initializing Whisper model: {model_name}")
        _WHISPER_MODEL = whisper.load_model(model_name)
    return _WHISPER_MODEL


def resolve_ffmpeg_path():
    ffmpeg_bin = shutil.which("ffmpeg")
    if ffmpeg_bin:
        return ffmpeg_bin

    home = os.path.expanduser("~")
    candidates = [
        os.path.join(home, "AppData", "Local", "Microsoft", "WinGet", "Links", "ffmpeg.exe"),
        os.path.join(home, "scoop", "shims", "ffmpeg.exe"),
        os.path.join(home, "AppData", "Local", "Programs", "ffmpeg", "bin", "ffmpeg.exe"),
        os.path.join(os.environ.get("ProgramFiles", "C:\\Program Files"), "ffmpeg", "bin", "ffmpeg.exe"),
        os.path.join(os.environ.get("ProgramFiles(x86)", "C:\\Program Files (x86)"), "ffmpeg", "bin", "ffmpeg.exe"),
    ]

    winget_pattern = os.path.join(
        home, "AppData", "Local", "Microsoft", "WinGet", "Packages",
        "*", "*", "bin", "ffmpeg.exe",
    )
    candidates.extend(glob.glob(winget_pattern))

    for candidate in candidates:
        if os.path.exists(candidate):
            ffmpeg_dir = os.path.dirname(candidate)
            os.environ["PATH"] = ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")
            return candidate

    return None


# ============================================================
# AUDIO RECORDING & TRANSCRIPTION
# ============================================================
def record_audio(duration=5, filename="mic_output.wav", stop_event=None):
    print("☥ Recording Started...")
    try:
        chunks = []

        def callback(indata, frames, time_info, status):
            if status:
                print(f"Audio stream status: {status}")
            chunks.append(indata.copy())

        with sd.InputStream(
            samplerate=samplerate, channels=1,
            dtype='float32', callback=callback,
        ):
            elapsed_ms = 0
            target_ms = int(duration * 1000)
            while elapsed_ms < target_ms:
                if stop_event is not None and stop_event.is_set():
                    print("⏹ Recording stopped by user")
                    return None
                sd.sleep(100)
                elapsed_ms += 100

        if not chunks:
            return None

        audio = np.concatenate(chunks, axis=0)
        write(filename, samplerate, (audio * 32767).astype('int16'))
        print("↪ Recording Saved as", filename)
        return filename
    except Exception as e:
        print("𒉽 Error while recording:", e)
        return None


def transcribe_audio(file_path):
    try:
        file_path = os.path.abspath(file_path)
        print(f"🎤 Starting transcription for: {file_path}")

        if not os.path.exists(file_path):
            raise RuntimeError(f"Audio file not found: {file_path}")

        ffmpeg_path = resolve_ffmpeg_path()
        if not ffmpeg_path:
            raise RuntimeError("FFmpeg binary not found. Please install ffmpeg.")

        try:
            import whisper.audio as whisper_audio
            whisper_audio.FFMPEG = ffmpeg_path
        except Exception:
            pass

        print(f"🎬 Using ffmpeg: {ffmpeg_path}")
        model = get_whisper_model()
        result = model.transcribe(file_path, fp16=False)
        return result["text"]
    except Exception as e:
        raise RuntimeError(f"Transcription failed: {str(e)}")


# ============================================================
# PDF EXTRACTION (unchanged)
# ============================================================
def sanitize_text_for_pdf(text):
    if not text:
        return text
    replacements = {
        '\u2013': '-', '\u2014': '--', '\u2018': "'", '\u2019': "'",
        '\u201C': '"', '\u201D': '"', '\u2022': '*', '\u2026': '...', '\xad': '',
    }
    for uc, repl in replacements.items():
        text = text.replace(uc, repl)
    try:
        text = text.encode('latin-1', errors='ignore').decode('latin-1')
    except Exception:
        pass
    return text


def extract_text_from_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip() if text.strip() else None
    except Exception:
        try:
            with open(file_path, 'rb') as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip() if text.strip() else None
        except Exception:
            return None


def extract_text_with_ocr(file_path):
    if not HAS_OCR:
        raise RuntimeError("OCR not available. Install pytesseract.")
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    image = page.to_image(resolution=300)
                    pil_image = image.original if hasattr(image, 'original') else Image.new('RGB', (image.width, image.height))
                    page_text = pytesseract.image_to_string(pil_image)
                    if page_text.strip():
                        text += f"[Page {page_num}]\n{page_text}\n"
                except Exception as page_error:
                    print(f"OCR failed on page {page_num}: {page_error}")
        return text.strip() if text.strip() else None
    except Exception as e:
        raise RuntimeError(f"OCR extraction failed: {str(e)}")


def process_pdf(file_path):
    try:
        text = extract_text_from_pdf(file_path)
        if text and len(text.strip()) > 50:
            return text
        ocr_text = extract_text_with_ocr(file_path)
        if ocr_text and len(ocr_text.strip()) > 50:
            return ocr_text
        raise RuntimeError("Could not extract text from PDF.")
    except Exception as e:
        raise RuntimeError(f"PDF processing failed: {str(e)}")


# ============================================================
# TEXT UTILITIES
# ============================================================
def chunk_text(text, max_words=1000):
    sentences = text.split('. ')
    chunks = []
    current_chunk = ""
    for sentence in sentences:
        if len(current_chunk.split()) + len(sentence.split()) < max_words:
            current_chunk += sentence + ". "
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks


def _truncate_for_prompt(text, max_chars=12000):
    """Truncate text to fit within prompt limits."""
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n[... content truncated for processing ...]"


def _split_sentences(text):
    if not text:
        return []
    normalized = text.replace("\r", "\n")
    normalized = re.sub(r'\n{2,}', '\n', normalized)
    parts = re.split(r'(?<=[.!?])\s+|\n+', normalized.strip())
    cleaned = [p.strip() for p in parts if len(p.strip()) > 20]
    return cleaned


def _sample_evenly(items, count):
    if not items:
        return []
    if len(items) <= count:
        return items
    if count <= 1:
        return [items[len(items) // 2]]

    step = (len(items) - 1) / (count - 1)
    idxs = []
    for i in range(count):
        idx = int(round(i * step))
        if idxs and idx <= idxs[-1]:
            idx = min(len(items) - 1, idxs[-1] + 1)
        idxs.append(idx)
    return [items[i] for i in idxs]


def _build_balanced_preview(text, max_chars):
    sentences = _split_sentences(text)
    if not sentences:
        return _truncate_for_prompt(text, max_chars)

    # Pull sentences from across the full transcript rather than only the beginning.
    candidate_count = min(len(sentences), max(20, max_chars // 220))
    sampled = _sample_evenly(sentences, candidate_count)

    out = []
    total = 0
    for sentence in sampled:
        line = sentence.strip()
        add_len = len(line) + 1
        if total + add_len > max_chars:
            break
        out.append(line)
        total += add_len
    return "\n".join(out)


def _guess_topic_title(text):
    sentences = _split_sentences(text)
    if not sentences:
        return "Lecture Topic"
    head = re.sub(r'(?i)section\s*\d+[:\-\s]*', '', sentences[0]).strip()
    head = re.sub(r'[^A-Za-z0-9\s-]', '', head).strip()
    words = head.split()
    return " ".join(words[:7]).title() if words else "Lecture Topic"


def _fallback_structured_summary(text, summary_length="Medium"):
    sentences = _split_sentences(text)
    level = (summary_length or "Medium").strip().lower()
    summary_count = 10 if level == "detailed" else 8
    key_count = 8 if level == "detailed" else 6
    important_count = 8 if level == "detailed" else 6
    breakdown_count = 6 if level == "detailed" else 4
    conclusion_count = 5 if level == "detailed" else 4

    if not sentences:
        return """### Summary: Lecture Topic
No usable transcript content was found.

---
### Key Concepts
- No key concepts detected.

---
### Main Topic Breakdown
- **Definition:** Please upload a clearer lecture file.
- **Structure:** Content could not be parsed.

---
### Summary of Important Points
- Better quality audio or text will produce richer notes.

---
### Conclusion
Try again with clearer audio or text-heavy material."""

    title = _guess_topic_title(text)
    summary_candidates = _sample_evenly(sentences, min(summary_count, len(sentences)))
    summary_block = " ".join(summary_candidates)

    key_candidates = _sample_evenly(sentences, min(key_count, len(sentences)))
    key_points = [f"- **Concept {i}:** {s}" for i, s in enumerate(key_candidates, 1)]

    important_candidates = _sample_evenly(sentences, min(important_count, len(sentences)))
    important_points = [f"- {s}" for s in important_candidates[:important_count]]

    conclusion_src = _sample_evenly(sentences[-max(8, len(sentences)//3):], min(conclusion_count, len(sentences)))
    conclusion_block = " ".join(conclusion_src)

    key_points_md = "\n".join(key_points)
    important_points_md = "\n".join(important_points)

    breakdown_src = _sample_evenly(sentences, min(breakdown_count, len(sentences)))
    labels = ["Definition", "Structure", "Key Detail", "Important Note", "Example", "Practical Note"]
    breakdown_lines = []
    for label, line in zip(labels, breakdown_src):
        breakdown_lines.append(f"- **{label}:** {line}")
    breakdown_md = "\n".join(breakdown_lines)

    return (
        f"### Summary: {title}\n"
        f"{summary_block}\n\n"
        "---\n"
        "### Key Concepts\n"
        f"{key_points_md}\n\n"
        "---\n"
        "### Main Topic Breakdown\n"
        f"{breakdown_md}\n\n"
        "---\n"
        "### Summary of Important Points\n"
        f"{important_points_md}\n\n"
        "---\n"
        "### Conclusion\n"
        f"{conclusion_block}"
    )


def _fallback_active_recall(text):
    def _question_from_sentence(sentence, idx):
        s = sentence.strip().rstrip('.!?')
        lower = s.lower()

        match = re.match(r'([A-Za-z][A-Za-z0-9\-\s]{2,50})\s+is\s+', s)
        if match:
            subject = match.group(1).strip()
            return f"What is {subject}?"

        if "used" in lower or "application" in lower:
            return "What are practical applications of this concept?"
        if "types" in lower or "kinds" in lower:
            return "What are the main types discussed?"
        if "important" in lower or "benefit" in lower:
            return "Why is this concept important?"

        words = s.split()
        snippet = " ".join(words[:8]).strip()
        return f"Explain this idea: {snippet}?"

    sentences = _split_sentences(text)
    cards = []
    for i, sentence in enumerate(sentences[:12], 1):
        cards.append({
            "question": _question_from_sentence(sentence, i),
            "answer": sentence,
        })

    if not cards:
        cards = [
            {
                "question": "What is the main topic of this lecture?",
                "answer": "The uploaded content did not contain enough text to extract robust flashcards.",
            },
            {
                "question": "How can better flashcards be generated?",
                "answer": "Use clearer audio or text-rich lecture material so key concepts can be extracted reliably.",
            },
        ]

    return {
        "qa_cards": cards,
        "mcqs": [],
        "true_false": [],
        "fill_blanks": [],
    }


# ============================================================
# SECTION 1: Concept Snapshot  (Gemini)
# ============================================================
def generate_concept_snapshot(text):
    """Generate a structured concept snapshot: what, why, where."""
    preview = _truncate_for_prompt(text, 10000)

    prompt = f"""You are an expert academic tutor. A student has provided a lecture transcript (or notes). 
Your job is to create a clear, synthesized "Concept Snapshot" that helps the student understand the topic quickly.

LECTURE CONTENT:
\"\"\"
{preview}
\"\"\"

Return a JSON object with exactly three keys:
{{
  "what": "A clear 4-6 sentence explanation of WHAT this topic is about. Identify the topic name, define it precisely, and explain the core idea. Do NOT echo transcript sentences — synthesize and explain in your own words like a textbook would.",
  "why": "A clear 3-5 sentence explanation of WHY this topic matters. Cover its importance in academics, exams, and foundational understanding for advanced topics.",
  "where": "A clear 3-5 sentence explanation of WHERE this topic is applied. Give concrete real-world examples and mention which fields/industries use it."
}}"""

    result = ollama_generate_json(prompt)
    if result and isinstance(result, dict):
        return {
            "what": result.get("what", ""),
            "why": result.get("why", ""),
            "where": result.get("where", ""),
        }
    # Fallback
    return {
        "what": "Could not generate snapshot.",
        "why": "Important for academic understanding.",
        "where": "Used across various domains.",
    }


# ============================================================
# SECTION 2: Core Concepts Breakdown  (Gemini)
# ============================================================
def generate_core_concepts(text):
    """Generate structured breakdown: definitions, formulas, key mechanisms, processes."""
    preview = _truncate_for_prompt(text, 12000)

    prompt = f"""You are an expert academic tutor creating structured study notes from a lecture.

LECTURE CONTENT:
\"\"\"
{preview}
\"\"\"

Extract and organize the core concepts into a JSON object with these keys:

{{
  "definitions": [
    "Clear, textbook-style definitions of key terms mentioned in the lecture. Each should be a complete sentence like: 'An array is a data structure that stores a collection of elements of the same data type in contiguous memory locations.' Write 5-12 definitions."
  ],
  "formulas": [
    "Any formulas, equations, or mathematical expressions mentioned. Write them clearly. If none exist, return an empty list."
  ],
  "mechanisms": [
    "Key mechanisms, principles, and how-it-works explanations. Each should be a clear, complete statement. Write 5-10 items."
  ],
  "processes": [
    "Step-by-step processes or methods described in the lecture. Each should be actionable and clear. Write 3-8 items."
  ]
}}

IMPORTANT: Do NOT copy transcript sentences verbatim. Synthesize and rephrase into clean, academic language."""

    result = ollama_generate_json(prompt)
    if result and isinstance(result, dict):
        return {
            "definitions": result.get("definitions", []),
            "formulas": result.get("formulas", []),
            "mechanisms": result.get("mechanisms", []),
            "processes": result.get("processes", []),
        }
    return {"definitions": [], "formulas": [], "mechanisms": [], "processes": []}


# ============================================================
# SECTION 3: Structured Summary  (Gemini)
# ============================================================
def generate_structured_summary(text, summary_length="Medium"):
    """Generate structured summary with markdown sections: Summary, Key Concepts, Important Points, Conclusion."""
    level = (summary_length or "Medium").strip().lower()
    if level == "detailed":
        preview = _build_balanced_preview(text, 18000)
        model_tokens = 1800
    elif level == "brief":
        preview = _build_balanced_preview(text, 8000)
        model_tokens = 700
    else:
        preview = _build_balanced_preview(text, 13000)
        model_tokens = 1200

    prompt = f"""You are an expert academic tutor creating professional study notes.

LECTURE CONTENT:
\"\"\"
{preview}
\"\"\"

Generate a structured summary in Markdown format that matches this style and section order:

### Summary: [topic title]
[A strong paragraph summary. Use 10-14 sentences in Detailed mode, 7-10 in Medium, 5-7 in Brief.]

---
### Key Concepts
[6-9 bullet points in Detailed mode, 5-7 in Medium, 4-5 in Brief. Start each with bold labels like "**Array Definition:** ..."]

---
### Main Topic Breakdown
[A detailed bullet breakdown with bold labels such as "**Definition:**", "**Structure:**", "**Analogy:**", "**Flexibility:**", "**Applications:**", "**Limitations:**"]

---
### Summary of Important Points
[6-8 concise bullet points in Detailed mode, 4-6 in Medium, 3-4 in Brief. Cover different parts of the lecture, not just opening and ending sections.]

---
### Conclusion
[A 4-6 sentence closing paragraph in Detailed mode, shorter in Medium/Brief.]

RETURN ONLY markdown content in this format. Do NOT add any extra section."""

    result = ollama_generate(prompt, temperature=0.3, max_tokens=model_tokens)
    required_sections = [
        "### Summary:",
        "### Key Concepts",
        "### Main Topic Breakdown",
        "### Summary of Important Points",
        "### Conclusion",
    ]
    if result and len(result) > 180 and all(section in result for section in required_sections):
        return result.strip()
    return _fallback_structured_summary(text, summary_length)


# ============================================================
# SECTION 4: Exam-Focused Insights  (Gemini) - DEPRECATED
# ============================================================
def generate_exam_insights(text):
    """Generate exam-focused content: FAQ, short/long answers, tricky areas."""
    preview = _truncate_for_prompt(text, 10000)

    prompt = f"""You are an expert exam coach preparing a student for their upcoming exam.

LECTURE CONTENT:
\"\"\"
{preview}
\"\"\"

Create exam-focused study material as a JSON object:

{{
  "faq_points": [
    "List 8-12 frequently asked exam questions or key points that are commonly tested. Each should be a clear, specific statement or question."
  ],
  "short_answers": [
    "Write 6-8 concise 2-mark style answers. Each should be a crisp 1-2 sentence answer to a likely exam question."
  ],
  "long_answers": [
    "Write 3-4 detailed 5-mark style answers. Each should be a well-structured paragraph covering the topic thoroughly."
  ],
  "tricky_areas": [
    "List 5-8 tricky or confusing aspects that students commonly get wrong in exams. Be specific about what the confusion is."
  ]
}}"""

    result = ollama_generate_json(prompt)
    if result and isinstance(result, dict):
        return {
            "faq_points": result.get("faq_points", []),
            "short_answers": result.get("short_answers", []),
            "long_answers": result.get("long_answers", []),
            "tricky_areas": result.get("tricky_areas", []),
        }
    return {"faq_points": [], "short_answers": [], "long_answers": [], "tricky_areas": []}


# ============================================================
# SECTION 4: Real-World Applications  (Gemini)
# ============================================================
def generate_real_world_applications(text):
    """Generate real-world applications and industry examples."""
    preview = _truncate_for_prompt(text, 8000)

    prompt = f"""You are an expert tutor helping students understand practical relevance.

LECTURE CONTENT:
\"\"\"
{preview}
\"\"\"

Identify the main topic from the lecture, then list 8-12 real-world applications.

Return a JSON object:
{{
  "applications": [
    "Each item should be a clear, specific real-world application with context. Example: 'Arrays are used in image processing where each pixel's RGB values are stored in a 2D array for manipulation.' Be concrete — mention industries, tools, or systems."
  ]
}}"""

    result = ollama_generate_json(prompt)
    if result and isinstance(result, dict):
        return result.get("applications", [])
    return []


# ============================================================
# SECTION 5: Common Mistakes & Confusions  (Gemini)
# ============================================================
def generate_common_mistakes(text):
    """Generate common mistakes and confusions."""
    preview = _truncate_for_prompt(text, 8000)

    prompt = f"""You are an experienced teacher who knows exactly where students go wrong.

LECTURE CONTENT:
\"\"\"
{preview}
\"\"\"

List 8-12 common mistakes, misconceptions, and confusions students have about this topic.

Return a JSON object:
{{
  "mistakes": [
    "Each item should clearly state the mistake AND the correct understanding. Example: 'Students often think arrays can store mixed data types (like integers and strings together), but in most languages, arrays require all elements to be of the same type.' Be specific and educational."
  ]
}}"""

    result = ollama_generate_json(prompt)
    if result and isinstance(result, dict):
        return result.get("mistakes", [])
    return []


# ============================================================
# SECTION 6: Smart Keywords + Definitions  (Gemini)
# ============================================================
def generate_smart_keywords(text):
    """Generate keywords with definitions and memory tricks."""
    preview = _truncate_for_prompt(text, 8000)

    prompt = f"""You are creating a keyword glossary with memory aids for a student.

LECTURE CONTENT:
\"\"\"
{preview}
\"\"\"

Identify 10-15 important keywords/terms from the lecture. For each, provide a clear definition and a memorable trick (mnemonic, analogy, or one-liner) to remember it.

Return a JSON object:
{{
  "keywords": [
    {{
      "term": "Array",
      "meaning": "A data structure that stores a fixed-size collection of elements of the same data type in contiguous memory locations.",
      "trick": "Think of an array like a row of lockers — each locker (index) holds one item, and all items must be the same type."
    }}
  ]
}}"""

    result = ollama_generate_json(prompt)
    if result and isinstance(result, dict):
        keywords = result.get("keywords", [])
        # Normalize keys
        normalized = []
        for kw in keywords:
            normalized.append({
                "term": kw.get("term", ""),
                "meaning": kw.get("meaning", ""),
                "trick": kw.get("trick", ""),
            })
        return normalized
    return []


# ============================================================
# SECTION 7: Active Recall Mode  (Gemini)
# ============================================================
def generate_active_recall(text):
    """Generate flashcards: Q&A, MCQs, True/False, Fill in the blanks."""
    preview = _truncate_for_prompt(text, 10000)

    prompt = f"""You are creating an active recall question bank for exam preparation.

LECTURE CONTENT:
\"\"\"
{preview}
\"\"\"

Generate a comprehensive set of practice questions in this JSON format:

{{
  "qa_cards": [
    {{
      "question": "What is an array?",
      "answer": "An array is a data structure that stores a fixed number of elements of the same data type in contiguous memory locations, allowing efficient access via indices."
    }}
  ],
  "mcqs": [
    {{
      "question": "Which of the following is true about arrays?",
      "options": ["Arrays can store mixed types", "Arrays store elements in contiguous memory", "Arrays have no fixed size", "Arrays are not data structures"],
      "answer": "Arrays store elements in contiguous memory"
    }}
  ],
  "true_false": [
    {{
      "statement": "An array can store elements of different data types simultaneously.",
      "answer": false,
      "explanation": "Arrays require all elements to be of the same data type. Storing mixed types is not allowed in traditional arrays."
    }}
  ],
  "fill_blanks": [
    {{
      "question": "An array stores elements in ________ memory locations.",
      "answer": "Contiguous"
    }}
  ]
}}

Generate:
- 8-10 Q&A flashcards covering key concepts
- 6-8 MCQs with 4 options each (only one correct)
- 5-6 True/False with explanations
- 6-8 Fill in the blanks

Make questions progressively harder. Ensure answers are accurate based on the lecture content."""

    result = ollama_generate_json(prompt, max_tokens=1200)
    if result and isinstance(result, dict):
        qa_cards = result.get("qa_cards", []) if isinstance(result.get("qa_cards", []), list) else []
        normalized_cards = []
        for i, card in enumerate(qa_cards, 1):
            if not isinstance(card, dict):
                continue
            q = str(card.get("question", "")).strip()
            a = str(card.get("answer", "")).strip()
            if not q or not a:
                continue
            # Replace generic placeholder questions with meaningful fallback wording.
            if re.search(r'key idea in point\s*\d+', q, flags=re.IGNORECASE):
                q = f"What is the key concept discussed in this part ({i})?"
            normalized_cards.append({"question": q, "answer": a})

        if normalized_cards:
            qa_cards = normalized_cards
        if qa_cards:
            return {
                "qa_cards": qa_cards,
                "mcqs": result.get("mcqs", []),
                "true_false": result.get("true_false", []),
                "fill_blanks": result.get("fill_blanks", []),
            }
    return _fallback_active_recall(text)


# ============================================================
# SECTION 8: Difficulty Level Indicator (rule-based, no model needed)
# ============================================================
def assess_difficulty_level(text):
    words = text.split()
    total_words = len(words)
    avg_word_len = sum(len(w) for w in words) / max(total_words, 1)
    long_words = sum(1 for w in words if len(w) > 8)
    long_ratio = long_words / max(total_words, 1)

    tech_patterns = [
        'algorithm', 'theorem', 'equation', 'derivative', 'integral',
        'hypothesis', 'methodology', 'paradigm', 'optimization', 'architecture',
        'computational', 'differential', 'polynomial', 'logarithmic', 'exponential',
        'regression', 'probability', 'distribution', 'correlation', 'inference',
        'quantum', 'molecular', 'electromagnetic', 'thermodynamic', 'entropy',
    ]
    tech_count = sum(1 for word in words if word.lower() in tech_patterns)

    sentences = re.split(r'[.!?]+', text)
    avg_sentence_len = total_words / max(len(sentences), 1)

    score = 0
    score += min(avg_word_len * 5, 25)
    score += min(long_ratio * 100, 25)
    score += min(tech_count * 3, 25)
    score += min(avg_sentence_len * 0.8, 25)

    if score < 33:
        level, label = "easy", "Easy"
        description = "Beginner-friendly content. Straightforward concepts that are easy to understand and memorize."
    elif score < 66:
        level, label = "moderate", "Moderate"
        description = "Intermediate level content. Requires careful reading and practice to fully understand."
    else:
        level, label = "advanced", "Advanced"
        description = "Complex material with technical depth. Multiple revisions recommended for exam preparation."

    return {
        "level": level, "label": label, "score": round(score),
        "description": description,
        "stats": {
            "total_words": total_words,
            "avg_word_length": round(avg_word_len, 1),
            "avg_sentence_length": round(avg_sentence_len, 1),
            "technical_terms": tech_count,
            "long_words_pct": round(long_ratio * 100, 1),
        },
    }


# ============================================================
# TOPIC EXTRACTION  (Gemini)
# ============================================================
def extract_main_topics(text):
    preview = _truncate_for_prompt(text, 6000)
    prompt = f"""Identify the 3-5 main topics covered in this lecture content. Return a JSON object:
{{
  "topics": ["Topic 1", "Topic 2", "Topic 3"]
}}

CONTENT:
\"\"\"
{preview}
\"\"\"
"""
    result = ollama_generate_json(prompt)
    if result and isinstance(result, dict):
        return result.get("topics", ["General Content"])
    return ["General Content"]


# ============================================================
# TEXT CLEANING  (Gemini)
# ============================================================
def clean_and_organize_text(text):
    """Use Ollama to clean and organize raw transcript text."""
    if not text or len(text) < 50:
        return text

    # Fast-path for long transcripts to avoid long local-model latency.
    if len(text) > 8000:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\b(um|uh|you know|like|basically)\b', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\s{2,}', ' ', text)
        return text.strip()

    text_to_clean = text

    prompt = f"""Clean and organize this raw lecture transcript into well-structured, readable text.

Rules:
- Fix grammar and sentence structure
- Remove filler words (um, uh, so basically, you know)
- Organize into clear paragraphs
- Keep ALL factual content — do not remove any information
- Do NOT add your own content — only restructure what's there
- Output plain text only (no markdown, no headers)

RAW TRANSCRIPT:
\"\"\"
{text_to_clean}
\"\"\"
"""

    cleaned = ollama_generate(prompt, temperature=0.1, max_tokens=800)
    if cleaned and len(cleaned) > 50:
        # If original was truncated, append remainder
        if len(text) > 15000:
            cleaned += "\n\n" + text[15000:]
        return cleaned
    # Fallback: basic regex cleaning
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)
    return text.strip()


# ============================================================
# EXPORT FUNCTIONS (PDF, Word, Markdown, JSON)
# ============================================================

def export_to_pdf(result_data):
    """Export comprehensive study notes to PDF"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', size=16)
        pdf.cell(200, 10, txt="Exam Study Notes", ln=True, align='C')
        pdf.set_font("Arial", size=10)
        pdf.cell(200, 5, txt="AI-Powered Smart Study Material", ln=True, align='C')
        pdf.ln(5)

        difficulty = result_data.get('difficulty', {})
        if difficulty:
            pdf.set_font("Arial", "B", size=10)
            pdf.cell(0, 6, f"Difficulty: {difficulty.get('label', 'N/A')} (Score: {difficulty.get('score', 'N/A')}/100)", ln=True)
            pdf.set_font("Arial", size=9)
            pdf.multi_cell(0, 5, sanitize_text_for_pdf(difficulty.get('description', '')))
            pdf.ln(3)

        # 1. Concept Snapshot
        snapshot = result_data.get('concept_snapshot', {})
        if snapshot:
            pdf.set_font("Arial", "B", size=13)
            pdf.cell(0, 8, "1. Concept Snapshot", ln=True)
            for label, key in [("What is this topic?", "what"), ("Why is it important?", "why"), ("Where is it used?", "where")]:
                pdf.set_font("Arial", "B", size=10)
                pdf.cell(0, 6, label, ln=True)
                pdf.set_font("Arial", size=9)
                pdf.multi_cell(0, 5, sanitize_text_for_pdf(snapshot.get(key, '')))
                pdf.ln(2)
            pdf.ln(3)

        # 2. Core Concepts
        core = result_data.get('core_concepts', {})
        if core:
            pdf.add_page()
            pdf.set_font("Arial", "B", size=13)
            pdf.cell(0, 8, "2. Core Concepts Breakdown", ln=True)
            for section_label, section_key in [("Definitions:", "definitions"), ("Important Formulas:", "formulas"),
                                                ("Key Mechanisms & Principles:", "mechanisms"), ("Step-by-Step Processes:", "processes")]:
                items = core.get(section_key, [])
                if items:
                    pdf.set_font("Arial", "B", size=10)
                    pdf.cell(0, 6, section_label, ln=True)
                    pdf.set_font("Arial", size=9)
                    for item in items:
                        pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"* {item}"))
                    pdf.ln(3)

        # 3. Exam-Focused Insights
        exam = result_data.get('exam_insights', {})
        if exam and any(exam.get(k) for k in ['faq_points', 'short_answers', 'long_answers', 'tricky_areas']):
            pdf.add_page()
            pdf.set_font("Arial", "B", size=13)
            pdf.cell(0, 8, "3. Exam-Focused Insights", ln=True)

            for section_label, section_key in [("Frequently Asked Points:", "faq_points"), ("2-Mark Answer Points:", "short_answers"),
                                                ("Tricky / Error-prone Areas:", "tricky_areas")]:
                items = exam.get(section_key, [])
                if items:
                    pdf.set_font("Arial", "B", size=10)
                    pdf.cell(0, 6, section_label, ln=True)
                    pdf.set_font("Arial", size=9)
                    for item in items:
                        pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"* {item}"))
                    pdf.ln(3)

            long_answers = exam.get('long_answers', [])
            if long_answers:
                pdf.set_font("Arial", "B", size=10)
                pdf.cell(0, 6, "5-Mark Answer Format:", ln=True)
                pdf.set_font("Arial", size=9)
                for i, la in enumerate(long_answers, 1):
                    pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"Answer {i}: {la}"))
                    pdf.ln(2)

        # 4. Real-World Applications
        apps = result_data.get('applications', [])
        if apps:
            pdf.add_page()
            pdf.set_font("Arial", "B", size=13)
            pdf.cell(0, 8, "4. Real-World Applications", ln=True)
            pdf.set_font("Arial", size=9)
            for app in apps:
                pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"* {app}"))
            pdf.ln(3)

        # 5. Common Mistakes
        mistakes = result_data.get('common_mistakes', [])
        if mistakes:
            pdf.set_font("Arial", "B", size=13)
            pdf.cell(0, 8, "5. Common Mistakes & Confusions", ln=True)
            pdf.set_font("Arial", size=9)
            for m in mistakes:
                pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"! {m}"))
            pdf.ln(3)

        # 6. Smart Keywords
        smart_kw = result_data.get('smart_keywords', [])
        if smart_kw:
            pdf.add_page()
            pdf.set_font("Arial", "B", size=13)
            pdf.cell(0, 8, "6. Smart Keywords + Definitions", ln=True)
            for kw in smart_kw:
                pdf.set_font("Arial", "B", size=10)
                pdf.cell(0, 6, sanitize_text_for_pdf(kw.get('term', '')), ln=True)
                pdf.set_font("Arial", size=9)
                pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"Meaning: {kw.get('meaning', '')}"))
                pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"Memory Trick: {kw.get('trick', '')}"))
                pdf.ln(2)

        # 7. Active Recall
        recall = result_data.get('active_recall', {})
        if recall:
            pdf.add_page()
            pdf.set_font("Arial", "B", size=13)
            pdf.cell(0, 8, "7. Active Recall Mode", ln=True)

            if recall.get('qa_cards'):
                pdf.set_font("Arial", "B", size=10)
                pdf.cell(0, 6, "Q&A Flashcards:", ln=True)
                for i, card in enumerate(recall['qa_cards'], 1):
                    pdf.set_font("Arial", "B", size=9)
                    pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"Q{i}: {card.get('question', '')}"))
                    pdf.set_font("Arial", size=9)
                    pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"A: {card.get('answer', '')}"))
                    pdf.ln(2)

            if recall.get('mcqs'):
                pdf.set_font("Arial", "B", size=10)
                pdf.cell(0, 6, "Multiple Choice Questions:", ln=True)
                for i, mcq in enumerate(recall['mcqs'], 1):
                    pdf.set_font("Arial", "B", size=9)
                    pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"Q{i}: {mcq.get('question', '')}"))
                    pdf.set_font("Arial", size=9)
                    for j, opt in enumerate(mcq.get('options', [])):
                        pdf.cell(0, 5, sanitize_text_for_pdf(f"  {chr(65+j)}) {opt}"), ln=True)
                    pdf.cell(0, 5, sanitize_text_for_pdf(f"  Answer: {mcq.get('answer', '')}"), ln=True)
                    pdf.ln(2)

            if recall.get('true_false'):
                pdf.set_font("Arial", "B", size=10)
                pdf.cell(0, 6, "True or False:", ln=True)
                for i, tf in enumerate(recall['true_false'], 1):
                    pdf.set_font("Arial", size=9)
                    ans_text = "True" if tf.get('answer') else "False"
                    pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"{i}. {tf.get('statement', '')}"))
                    pdf.cell(0, 5, sanitize_text_for_pdf(f"   Answer: {ans_text} - {tf.get('explanation', '')}"), ln=True)
                    pdf.ln(1)

            if recall.get('fill_blanks'):
                pdf.set_font("Arial", "B", size=10)
                pdf.cell(0, 6, "Fill in the Blanks:", ln=True)
                for i, fb in enumerate(recall['fill_blanks'], 1):
                    pdf.set_font("Arial", size=9)
                    pdf.multi_cell(0, 5, sanitize_text_for_pdf(f"{i}. {fb.get('question', '')}"))
                    pdf.cell(0, 5, sanitize_text_for_pdf(f"   Answer: {fb.get('answer', '')}"), ln=True)
                    pdf.ln(1)

        pdf.output("exam_study_notes.pdf")
        print("PDF exported successfully")
        return "exam_study_notes.pdf"
    except Exception as e:
        raise RuntimeError(f"PDF export failed: {str(e)}")


def export_to_markdown(result_data):
    """Export study notes to Markdown format"""
    try:
        md = "# Exam Study Notes\n\n"
        md += "*AI-Powered Smart Study Material*\n\n---\n\n"

        difficulty = result_data.get('difficulty', {})
        if difficulty:
            icons = {"easy": "🟢", "moderate": "🟡", "advanced": "🔴"}
            icon = icons.get(difficulty.get('level', ''), '⚪')
            md += f"## 📊 Difficulty: {icon} {difficulty.get('label', 'N/A')} (Score: {difficulty.get('score', 'N/A')}/100)\n\n"
            md += f"{difficulty.get('description', '')}\n\n---\n\n"

        snapshot = result_data.get('concept_snapshot', {})
        if snapshot:
            md += "## 🧠 1. Concept Snapshot\n\n"
            md += f"**What is this topic?**\n{snapshot.get('what', '')}\n\n"
            md += f"**Why is it important?**\n{snapshot.get('why', '')}\n\n"
            md += f"**Where is it used?**\n{snapshot.get('where', '')}\n\n---\n\n"

        core = result_data.get('core_concepts', {})
        if core:
            md += "## 🏗️ 2. Core Concepts Breakdown\n\n"
            for label, key in [("### Definitions\n", "definitions"), ("### Important Formulas\n", "formulas"),
                               ("### Key Mechanisms & Principles\n", "mechanisms"), ("### Step-by-Step Processes\n", "processes")]:
                items = core.get(key, [])
                if items:
                    md += label
                    for item in items:
                        md += f"- {item}\n"
                    md += "\n"
            md += "---\n\n"

        exam = result_data.get('exam_insights', {})
        if exam:
            md += "## 🎯 3. Exam-Focused Insights\n\n"
            for label, key in [("### Frequently Asked Points\n", "faq_points"), ("### 2-Mark Answer Points\n", "short_answers")]:
                items = exam.get(key, [])
                if items:
                    md += label
                    for item in items:
                        md += f"- {item}\n"
                    md += "\n"
            long_answers = exam.get('long_answers', [])
            if long_answers:
                md += "### 5-Mark Answer Format\n"
                for i, la in enumerate(long_answers, 1):
                    md += f"**Answer {i}:** {la}\n\n"
            tricky = exam.get('tricky_areas', [])
            if tricky:
                md += "### ⚠️ Tricky Areas\n"
                for t in tricky:
                    md += f"- {t}\n"
                md += "\n"
            md += "---\n\n"

        apps = result_data.get('applications', [])
        if apps:
            md += "## 🧩 4. Real-World Applications\n\n"
            for app in apps:
                md += f"- {app}\n"
            md += "\n---\n\n"

        mistakes = result_data.get('common_mistakes', [])
        if mistakes:
            md += "## ⚠️ 5. Common Mistakes & Confusions\n\n"
            for m in mistakes:
                md += f"- {m}\n"
            md += "\n---\n\n"

        smart_kw = result_data.get('smart_keywords', [])
        if smart_kw:
            md += "## 🔑 6. Smart Keywords + Definitions\n\n"
            md += "| Term | Meaning | Memory Trick |\n"
            md += "|------|---------|-------------|\n"
            for kw in smart_kw:
                md += f"| **{kw.get('term', '')}** | {kw.get('meaning', '')[:80]} | {kw.get('trick', '')} |\n"
            md += "\n---\n\n"

        recall = result_data.get('active_recall', {})
        if recall:
            md += "## 🃏 7. Active Recall Mode\n\n"
            if recall.get('qa_cards'):
                md += "### Q&A Flashcards\n"
                for i, card in enumerate(recall['qa_cards'], 1):
                    md += f"**Q{i}:** {card.get('question', '')}\n"
                    md += f"**A:** {card.get('answer', '')}\n\n"
            if recall.get('mcqs'):
                md += "### MCQs\n"
                for i, mcq in enumerate(recall['mcqs'], 1):
                    md += f"**Q{i}:** {mcq.get('question', '')}\n"
                    for j, opt in enumerate(mcq.get('options', [])):
                        md += f"  {chr(65+j)}) {opt}\n"
                    md += f"  **Answer:** {mcq.get('answer', '')}\n\n"
            if recall.get('true_false'):
                md += "### True or False\n"
                for i, tf in enumerate(recall['true_false'], 1):
                    ans = "True" if tf.get('answer') else "False"
                    md += f"{i}. {tf.get('statement', '')} — **{ans}**\n"
                    md += f"   *{tf.get('explanation', '')}*\n"
                md += "\n"
            if recall.get('fill_blanks'):
                md += "### Fill in the Blanks\n"
                for i, fb in enumerate(recall['fill_blanks'], 1):
                    md += f"{i}. {fb.get('question', '')}\n   **Answer:** {fb.get('answer', '')}\n\n"
            md += "---\n\n"

        md += f"\n*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n"

        with open("exam_study_notes.md", "w", encoding="utf-8") as f:
            f.write(md)
        print("Markdown exported successfully")
        return "exam_study_notes.md"
    except Exception as e:
        raise RuntimeError(f"Markdown export failed: {str(e)}")


def export_to_word(result_data):
    """Export comprehensive study notes to Word"""
    try:
        doc = Document()
        title = doc.add_heading("Exam Study Notes", 0)
        title.alignment = 1
        doc.add_paragraph("AI-Powered Smart Study Material").alignment = 1

        difficulty = result_data.get('difficulty', {})
        if difficulty:
            doc.add_heading(f"Difficulty: {difficulty.get('label', 'N/A')} (Score: {difficulty.get('score', 'N/A')}/100)", level=2)
            doc.add_paragraph(difficulty.get('description', ''))

        snapshot = result_data.get('concept_snapshot', {})
        if snapshot:
            doc.add_heading("1. Concept Snapshot", level=1)
            for label, key in [("What is this topic?", "what"), ("Why is it important?", "why"), ("Where is it used?", "where")]:
                doc.add_heading(label, level=2)
                doc.add_paragraph(snapshot.get(key, ''))

        core = result_data.get('core_concepts', {})
        if core:
            doc.add_page_break()
            doc.add_heading("2. Core Concepts Breakdown", level=1)
            for label, key in [("Definitions", "definitions"), ("Important Formulas", "formulas"),
                               ("Key Mechanisms & Principles", "mechanisms"), ("Step-by-Step Processes", "processes")]:
                items = core.get(key, [])
                if items:
                    doc.add_heading(label, level=2)
                    for item in items:
                        doc.add_paragraph(item, style='List Bullet')

        exam = result_data.get('exam_insights', {})
        if exam and any(exam.get(k) for k in ['faq_points', 'short_answers', 'long_answers', 'tricky_areas']):
            doc.add_page_break()
            doc.add_heading("3. Exam-Focused Insights", level=1)
            for label, key in [("Frequently Asked Points", "faq_points"), ("2-Mark Answer Points", "short_answers"),
                               ("Tricky Areas", "tricky_areas")]:
                items = exam.get(key, [])
                if items:
                    doc.add_heading(label, level=2)
                    for item in items:
                        doc.add_paragraph(item, style='List Bullet')
            long_answers = exam.get('long_answers', [])
            if long_answers:
                doc.add_heading("5-Mark Answer Format", level=2)
                for i, la in enumerate(long_answers, 1):
                    doc.add_paragraph(f"Answer {i}: {la}")

        apps = result_data.get('applications', [])
        if apps:
            doc.add_page_break()
            doc.add_heading("4. Real-World Applications", level=1)
            for app in apps:
                doc.add_paragraph(app, style='List Bullet')

        mistakes = result_data.get('common_mistakes', [])
        if mistakes:
            doc.add_heading("5. Common Mistakes & Confusions", level=1)
            for m in mistakes:
                doc.add_paragraph(m, style='List Bullet')

        smart_kw = result_data.get('smart_keywords', [])
        if smart_kw:
            doc.add_page_break()
            doc.add_heading("6. Smart Keywords + Definitions", level=1)
            for kw in smart_kw:
                doc.add_heading(kw.get('term', ''), level=2)
                doc.add_paragraph(f"Meaning: {kw.get('meaning', '')}")
                doc.add_paragraph(f"Memory Trick: {kw.get('trick', '')}")

        recall = result_data.get('active_recall', {})
        if recall:
            doc.add_page_break()
            doc.add_heading("7. Active Recall Mode", level=1)
            if recall.get('qa_cards'):
                doc.add_heading("Q&A Flashcards", level=2)
                for i, card in enumerate(recall['qa_cards'], 1):
                    doc.add_paragraph(f"Q{i}: {card.get('question', '')}")
                    doc.add_paragraph(f"A: {card.get('answer', '')}")
            if recall.get('mcqs'):
                doc.add_heading("MCQs", level=2)
                for i, mcq in enumerate(recall['mcqs'], 1):
                    doc.add_paragraph(f"Q{i}: {mcq.get('question', '')}")
                    for j, opt in enumerate(mcq.get('options', [])):
                        doc.add_paragraph(f"  {chr(65+j)}) {opt}")
                    doc.add_paragraph(f"  Answer: {mcq.get('answer', '')}")
            if recall.get('true_false'):
                doc.add_heading("True or False", level=2)
                for i, tf in enumerate(recall['true_false'], 1):
                    ans = "True" if tf.get('answer') else "False"
                    doc.add_paragraph(f"{i}. {tf.get('statement', '')} -- {ans}")
            if recall.get('fill_blanks'):
                doc.add_heading("Fill in the Blanks", level=2)
                for i, fb in enumerate(recall['fill_blanks'], 1):
                    doc.add_paragraph(f"{i}. {fb.get('question', '')}")
                    doc.add_paragraph(f"   Answer: {fb.get('answer', '')}")

        doc.save("exam_study_notes.docx")
        print("Word exported successfully")
        return "exam_study_notes.docx"
    except Exception as e:
        raise RuntimeError(f"Word export failed: {str(e)}")


def export_to_json(result_data):
    try:
        with open("lecture_summary.json", "w", encoding="utf-8") as f:
            json.dump(result_data, f, ensure_ascii=False, indent=2, default=str)
        print("JSON exported successfully")
    except Exception as e:
        raise RuntimeError(f"JSON export failed: {str(e)}")


# ============================================================
# SUMMARY LENGTH PROFILES
# ============================================================
def _limit_text(value, max_chars):
    if not isinstance(value, str):
        return value
    return value[:max_chars].rstrip() + "..." if len(value) > max_chars else value


def _limit_list(value, count):
    return value[:count] if isinstance(value, list) else value


def apply_summary_length_profile(result_data, summary_length="Medium"):
    level = (summary_length or "Medium").strip().lower()
    if level == "detailed":
        result_data["summary_length"] = "Detailed"
        return result_data

    if level == "brief":
        list_limits = {
            "topics": 5, "applications": 4, "common_mistakes": 4, "smart_keywords": 8,
            "core_concepts": {"definitions": 3, "formulas": 2, "mechanisms": 3, "processes": 2},
            "exam_insights": {"faq_points": 4, "short_answers": 4, "long_answers": 2, "tricky_areas": 3},
            "active_recall": {"qa_cards": 6, "mcqs": 4, "true_false": 4, "fill_blanks": 4},
        }
        text_limits = {"cleaned_text": 4000, "concept_snapshot": {"what": 320, "why": 220, "where": 220}}
        label = "Brief"
    else:
        list_limits = {
            "topics": 12, "applications": 10, "common_mistakes": 10, "smart_keywords": 15,
            "core_concepts": {"definitions": 10, "formulas": 6, "mechanisms": 10, "processes": 8},
            "exam_insights": {"faq_points": 10, "short_answers": 10, "long_answers": 5, "tricky_areas": 8},
            "active_recall": {"qa_cards": 12, "mcqs": 8, "true_false": 8, "fill_blanks": 8},
        }
        text_limits = {"cleaned_text": 18000, "concept_snapshot": {"what": 1000, "why": 700, "where": 700}}
        label = "Medium"

    for key in ["topics", "applications", "common_mistakes", "smart_keywords"]:
        result_data[key] = _limit_list(result_data.get(key, []), list_limits[key])

    for section_key in ["core_concepts", "exam_insights", "active_recall"]:
        section = result_data.get(section_key, {})
        for k, count in list_limits[section_key].items():
            section[k] = _limit_list(section.get(k, []), count)
        result_data[section_key] = section

    snapshot = result_data.get("concept_snapshot", {})
    for k in ["what", "why", "where"]:
        snapshot[k] = _limit_text(snapshot.get(k, ""), text_limits["concept_snapshot"][k])
    result_data["concept_snapshot"] = snapshot
    result_data["cleaned_text"] = _limit_text(result_data.get("cleaned_text", ""), text_limits["cleaned_text"])
    result_data["summary_length"] = label
    return result_data


# ============================================================
# MAIN PROCESSING PIPELINE
# ============================================================
def process_input(source_type="mic", file_path=None, duration=10, export_format="PDF",
                  pdf_text=None, stop_event=None, summary_length="Medium"):
    try:
        # Check if Ollama is running
        if not get_ollama_client():
            return {"error": "Ollama is not running. Install Ollama from https://ollama.ai and run 'ollama serve' in a terminal, then 'ollama pull mistral' in another terminal."}
        
        cleanup_files = []

        if source_type == "mic":
            file_path = record_audio(duration=duration, stop_event=stop_event)
            if not file_path:
                if stop_event is not None and stop_event.is_set():
                    return {"error": "Recording stopped by user."}
                return {"error": "Microphone recording failed."}
            transcript = transcribe_audio(file_path)
            cleanup_files.append(file_path)
        elif source_type == "file" and file_path:
            transcript = transcribe_audio(file_path)
        elif source_type == "pdf" and pdf_text:
            transcript = pdf_text
        else:
            return {"error": "Invalid input source."}

        print("📝 Cleaning and organizing text...")
        cleaned_text = clean_and_organize_text(transcript)

        print("🔍 Identifying main topics...")
        topics = extract_main_topics(cleaned_text)

        print("1/4 Generating Concept Snapshot...")
        concept_snapshot = generate_concept_snapshot(cleaned_text)

        print("2/4 Building Core Concepts Breakdown...")
        core_concepts = generate_core_concepts(cleaned_text)

        print("3/4 Generating Structured Summary...")
        structured_summary = generate_structured_summary(cleaned_text, summary_length)

        print("4/4 Building Active Recall Questions...")
        active_recall = generate_active_recall(cleaned_text)

        result_data = {
            "concept_snapshot": concept_snapshot,
            "core_concepts": core_concepts,
            "structured_summary": structured_summary,
            "active_recall": active_recall,
            "topics": topics,
            "cleaned_text": cleaned_text,
        }

        result_data = apply_summary_length_profile(result_data, summary_length)

        output_file = ""
        export_format = export_format.upper()

        if export_format == "PDF":
            output_file = export_to_pdf(result_data)
        elif export_format == "WORD":
            output_file = export_to_word(result_data)
        elif export_format == "MARKDOWN":
            output_file = export_to_markdown(result_data)
        elif export_format == "JSON":
            export_to_json(result_data)
            output_file = "lecture_summary.json"
        else:
            return {"error": "Unsupported export format."}

        for f in cleanup_files:
            if os.path.exists(f):
                os.remove(f)

        result_data["output_file"] = output_file
        return result_data

    except Exception as e:
        return {"error": str(e)}


if __name__ == "__main__":
    print("֎ Choose input method:")
    print("1.  Mic")
    print("2.  File")
    choice = input("Enter choice [1/2]: ")

    if choice == "1":
        duration = float(input(" Enter duration (minutes): "))
        result = process_input(source_type="mic", duration=int(duration * 60), export_format="PDF")
    elif choice == "2":
        path = input(" Enter file path: ")
        result = process_input(source_type="file", file_path=path, export_format="PDF")
    else:
        print("➶ Invalid choice.")
        exit()

    if "error" in result:
        print(" Error:", result["error"])
    else:
        print("𒉽 Summary generated! File saved as:", result["output_file"])